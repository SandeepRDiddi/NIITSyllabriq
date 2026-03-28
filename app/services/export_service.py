from __future__ import annotations

import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app.core.config import settings

# ---------------------------------------------------------------------------
# NIIT StackRoute brand palette (from the official design document template)
# ---------------------------------------------------------------------------
SR_ORANGE = RGBColor(0xE9, 0x71, 0x32)   # #E97132 — cover title
SR_NAVY   = RGBColor(0x0F, 0x47, 0x61)   # #0F4761 — section headings
SR_ACCENT = RGBColor(0x5B, 0x9B, 0xD4)   # #5B9BD4 — accent / sub-labels
SR_BLACK  = RGBColor(0x00, 0x00, 0x00)   # #000000 — body text
SR_GREY   = RGBColor(0x6B, 0x7A, 0x99)   # #6B7A99 — footer / disclaimer
SR_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)

# Fixed boilerplate — must appear verbatim as per the official template
_LEARNING_PEDAGOGY_TEXT = (
    "The pedagogic model is focused on experiential learning (In person and remote virtual learning) "
    "mode. Some expert mentors shall work with students through the program. Learning is in an "
    "environment that combines the convenience of anytime access with the intensity of mentoring.\n\n"
    "The model combines the following elements:\n\n"
    "1. Instructor-led Live connects: These work on a fixed schedule with recorded versions available "
    "to people who missed them.\n"
    "   - Sessions that provide context.\n"
    "   - Sessions that demonstrate the usage of tools or technologies\n"
    "   - Sessions with expert-led demonstrations that provide step-by-step guidance on critical tasks.\n"
    "   - Sessions that explain best practices.\n"
    "   - Sessions that explain common pitfalls/issues.\n"
    "   - Sessions that discuss success stories, case studies and real-world scenarios that provide "
    "insight into the practical challenges and solutions.\n"
    "2. Reference learning material."
)

_ABOUT_STACKROUTE_TEXT = (
    "Established in August 2015, StackRoute\u00ae is an NIIT incubated venture. StackRoute provides "
    "disruptive IT Learning solutions that produce top-class full-stack developers & tech professionals "
    "with deep skills. We have evolved a mechanism of providing immersive experiences backed by mastery "
    "learning and individual tutoring that allows us to guarantee outcomes. As a digital transformation "
    "partner, StackRoute works with several large, mid & small global IT organizations, Global Incubation "
    "Centers (GICs), Global Capability Centers (GCCs) & product engineering teams."
)

_IP_NOTICE = (
    "All Information within this document is Intellectual property of NIIT (NIIT Ltd). "
    "No part of this document or the program design or program structure mentioned within "
    "can be shared or used within any organization without the permission of NIIT (NIIT Ltd)."
)

# Sections that carry fixed boilerplate — skip them in the markdown body pass
_FIXED_SECTIONS = {"learning pedagogy", "about stackroute"}


def _set_cell_bg(cell, hex_color: str) -> None:
    """Set table cell background colour via OOXML."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_font(run, size_pt: float, bold: bool = False, italic: bool = False, color: RGBColor | None = None, font_name: str = "Candara") -> None:
    """Apply consistent StackRoute font settings to a run."""
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


class ExportService:

    def export_docx(self, title: str, content: str, stem: str) -> str:
        target = Path(settings.export_dir) / f"{stem}.docx"
        doc = Document()

        # ── Page layout (A4, 1-inch margins) ─────────────────────────────────
        for section in doc.sections:
            section.top_margin    = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin   = Inches(1.0)
            section.right_margin  = Inches(1.0)

        # ── Cover page ────────────────────────────────────────────────────────
        self._add_cover(doc, content)

        doc.add_paragraph("")

        # ── Body content ──────────────────────────────────────────────────────
        self._render_body(doc, content)

        # ── Fixed boilerplate sections ────────────────────────────────────────
        # These are appended last because they are template-mandated and
        # must appear verbatim regardless of what the LLM generated.
        self._add_fixed_section(doc, "Learning Pedagogy", _LEARNING_PEDAGOGY_TEXT, size_pt=14)
        self._add_fixed_section(doc, "About StackRoute", _ABOUT_STACKROUTE_TEXT, size_pt=14, italic_body=True)

        # Note: sign-off status is tracked in the database (DesignDocument.status).
        # No approval table is printed in the document itself.

        # ── Footer notice (inline at end of doc) ───────────────────────────────
        self._add_footer_notice(doc)

        doc.save(target)
        return str(target)

    # ------------------------------------------------------------------
    # Cover page
    # ------------------------------------------------------------------

    def _add_cover(self, doc: Document, content: str) -> None:
        """
        Build the cover block.
        Extracts program_name and total_duration from the markdown content,
        matching the official NIIT StackRoute cover layout.
        """
        # Extract course title: "# Course: <name>" or "Course: <name>"
        program_name = self._extract_cover_field(content, r"#\s*Course:\s*(.+)")
        if not program_name:
            program_name = self._extract_cover_field(content, r"Course:\s*(.+)")
        if not program_name:
            program_name = "[Program Name]"

        # Extract total duration
        duration = self._extract_cover_field(content, r"\*\*Total Duration:\*\*\s*(.+)")
        if not duration:
            duration = self._extract_cover_field(content, r"Total Duration:\s*(.+)")

        # NIIT StackRoute wordmark
        brand_para = doc.add_paragraph()
        brand_run = brand_para.add_run("NIIT StackRoute®")
        _set_font(brand_run, 16, bold=True, color=SR_NAVY)
        brand_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        doc.add_paragraph("")

        # "Course: <Program Name>" — 28pt bold orange (exact template spec)
        title_para = doc.add_paragraph()
        label_run = title_para.add_run("Course: ")
        _set_font(label_run, 28, bold=True, color=SR_ORANGE)
        name_run = title_para.add_run(program_name)
        _set_font(name_run, 28, bold=True, color=SR_ORANGE)
        title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        if duration:
            doc.add_paragraph("")
            dur_para = doc.add_paragraph()
            dur_label = dur_para.add_run("Total Duration: ")
            _set_font(dur_label, 12, bold=True, color=SR_BLACK)
            dur_val = dur_para.add_run(duration)
            _set_font(dur_val, 12, bold=False, color=SR_BLACK)

        doc.add_paragraph("")

        # Metadata strip
        meta = doc.add_table(rows=3, cols=2)
        meta.style = "Table Grid"
        meta_rows = [
            ("Document Type",  "Program Design Document"),
            ("Prepared by",    "NIIT StackRoute — Design Automation"),
            ("Date",           date.today().strftime("%d %B %Y")),
        ]
        for i, (lbl, val) in enumerate(meta_rows):
            lc = meta.rows[i].cells[0]
            vc = meta.rows[i].cells[1]
            _set_cell_bg(lc, "0F4761")
            lbl_run = lc.paragraphs[0].add_run(lbl)
            _set_font(lbl_run, 9, bold=True, color=SR_WHITE)
            val_run = vc.paragraphs[0].add_run(val)
            _set_font(val_run, 9, color=SR_BLACK)

        doc.add_paragraph("")
        # Decorative divider
        hr = doc.add_paragraph()
        hr_run = hr.add_run("─" * 80)
        _set_font(hr_run, 7, color=SR_NAVY)

    # ------------------------------------------------------------------
    # Body rendering
    # ------------------------------------------------------------------

    def _render_body(self, doc: Document, content: str) -> None:
        """
        Parse the markdown body and render into Word paragraphs.
        Skips the YAML front-matter, the cover title, duration field, and
        the fixed boilerplate sections (Learning Pedagogy, About StackRoute)
        which are added separately with guaranteed correct content.
        """
        lines = content.splitlines()
        i = 0
        in_fixed_section = False
        skip_approval = False
        yaml_consumed = False  # Only strip the front-matter block once

        while i < len(lines):
            line = lines[i]

            # ── Skip YAML front-matter (first ---...--- block only) ───────────
            if line.strip() == "---" and not yaml_consumed:
                i += 1
                while i < len(lines) and lines[i].strip() != "---":
                    i += 1
                i += 1          # skip closing ---
                yaml_consumed = True
                continue

            # ── All subsequent --- are horizontal rules — render as whitespace ─
            if line.strip() == "---":
                doc.add_paragraph("")
                i += 1
                continue

            # ── Skip cover-level H1 and Total Duration ────────────────────────
            if line.startswith("# ") or re.match(r"\*\*Total Duration", line):
                i += 1
                continue

            # ── H2 section headings ───────────────────────────────────────────
            if line.startswith("## "):
                heading_text = line[3:].strip()
                heading_lower = heading_text.lower()

                # Track whether we're inside a fixed boilerplate section
                if heading_lower in _FIXED_SECTIONS:
                    in_fixed_section = True
                    i += 1
                    continue
                elif any(heading_lower.startswith(k) for k in ("approval", "sign-off", "sign off")):
                    skip_approval = True
                    i += 1
                    continue
                else:
                    in_fixed_section = False
                    skip_approval = False

                # Determine size: major sections (Program Introduction,
                # Indicative Design) get 16pt; all others get 14pt
                major = heading_lower in (
                    "program introduction",
                    "indicative design and content coverage",
                )
                size = 16 if major else 14

                para = doc.add_heading(heading_text, level=1)
                for run in para.runs:
                    _set_font(run, size, bold=False, color=SR_NAVY)
                i += 1
                continue

            # ── Skip content that belongs to fixed/approval sections ──────────
            if in_fixed_section or skip_approval:
                i += 1
                continue

            # ── H3 sub-headings ───────────────────────────────────────────────
            if line.startswith("### "):
                para = doc.add_heading(line[4:].strip(), level=2)
                for run in para.runs:
                    _set_font(run, 12, bold=True, color=SR_NAVY)
                i += 1
                continue

            # ── Markdown table ────────────────────────────────────────────────
            # Use lstrip() so lines like " | col |" (LLM adds leading space) are caught
            if line.lstrip().startswith("|"):
                table_lines: list[str] = []
                while i < len(lines) and lines[i].lstrip().startswith("|"):
                    table_lines.append(lines[i].strip())  # normalise leading/trailing whitespace
                    i += 1
                rows = [r for r in table_lines if not re.match(r"^\|[-| :]+\|$", r.strip())]
                if rows:
                    cols = [c.strip() for c in rows[0].strip("|").split("|")]
                    tbl = doc.add_table(rows=len(rows), cols=len(cols))
                    tbl.style = "Table Grid"
                    for ri, row_line in enumerate(rows):
                        cells = [c.strip() for c in row_line.strip("|").split("|")]
                        for ci, cell_text in enumerate(cells):
                            if ci < len(tbl.rows[ri].cells):
                                cell = tbl.rows[ri].cells[ci]
                                if ri == 0:
                                    _set_cell_bg(cell, "0F4761")
                                    run = cell.paragraphs[0].add_run(cell_text)
                                    _set_font(run, 9, bold=True, color=SR_WHITE)
                                else:
                                    run = cell.paragraphs[0].add_run(cell_text)
                                    _set_font(run, 9, color=SR_BLACK)
                    doc.add_paragraph("")
                continue

            # ── Bullet / list item ────────────────────────────────────────────
            if re.match(r"^[-*]\s", line) or re.match(r"^\s+[-*]\s", line):
                # Detect indented sub-bullet
                indent = len(line) - len(line.lstrip())
                style = "List Bullet 2" if indent >= 3 else "List Bullet"
                text = re.sub(r"^[\s\-*]+", "", line).strip()
                para = doc.add_paragraph(style=style)
                run = para.add_run(text)
                _set_font(run, 10, color=SR_BLACK)
                i += 1
                continue

            # ── Numbered list item ────────────────────────────────────────────
            if re.match(r"^\d+\.", line.strip()):
                para = doc.add_paragraph(style="List Number")
                run = para.add_run(re.sub(r"^\d+\.\s*", "", line.strip()))
                _set_font(run, 10, color=SR_BLACK)
                i += 1
                continue

            # ── Horizontal rule ───────────────────────────────────────────────
            if line.strip() == "---":
                doc.add_paragraph("")
                i += 1
                continue

            # ── Bold / italic inline handling + regular paragraph ─────────────
            stripped = line.strip()
            if stripped:
                para = doc.add_paragraph()
                self._add_inline_text(para, stripped)
            i += 1

    def _add_inline_text(self, para, text: str) -> None:
        """Handle **bold** and *italic* markers within a line."""
        parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = para.add_run(part[2:-2])
                _set_font(run, 10, bold=True, color=SR_BLACK)
            elif part.startswith("*") and part.endswith("*"):
                run = para.add_run(part[1:-1])
                _set_font(run, 10, italic=True, color=SR_BLACK)
            elif part:
                run = para.add_run(part)
                _set_font(run, 10, color=SR_BLACK)

    # ------------------------------------------------------------------
    # Fixed boilerplate sections
    # ------------------------------------------------------------------

    def _add_fixed_section(
        self,
        doc: Document,
        heading: str,
        body_text: str,
        size_pt: float = 14,
        italic_body: bool = False,
    ) -> None:
        """
        Add a mandatory fixed section with its exact boilerplate text.
        This ensures the output always contains the correct content
        regardless of what the LLM generated.
        """
        doc.add_paragraph("")
        para = doc.add_heading(heading, level=1)
        for run in para.runs:
            _set_font(run, size_pt, bold=False, color=SR_NAVY)

        for block in body_text.split("\n"):
            stripped = block.strip()
            if not stripped:
                doc.add_paragraph("")
                continue
            if re.match(r"^\d+\.", stripped):
                p = doc.add_paragraph(style="List Number")
                run = p.add_run(re.sub(r"^\d+\.\s*", "", stripped))
                _set_font(run, 10, italic=italic_body, color=SR_BLACK)
            elif re.match(r"^[-*]\s", stripped):
                indent = len(block) - len(block.lstrip())
                style = "List Bullet 2" if indent >= 3 else "List Bullet"
                p = doc.add_paragraph(style=style)
                run = p.add_run(re.sub(r"^[-*]\s+", "", stripped))
                _set_font(run, 10, italic=italic_body, color=SR_BLACK)
            else:
                p = doc.add_paragraph()
                run = p.add_run(stripped)
                _set_font(run, 10, italic=italic_body, color=SR_BLACK)

    # ------------------------------------------------------------------
    # Approval table and footer
    # ------------------------------------------------------------------

    def _add_footer_notice(self, doc: Document) -> None:
        doc.add_paragraph("")
        # Copyright line
        cp_para = doc.add_paragraph()
        cp_run = cp_para.add_run(f"© NIIT {date.today().year}")
        _set_font(cp_run, 8, color=SR_GREY)
        cp_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # IP Notice
        ip_para = doc.add_paragraph()
        ip_run = ip_para.add_run(_IP_NOTICE)
        _set_font(ip_run, 8, italic=True, color=SR_GREY)
        ip_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------

    def _extract_cover_field(self, content: str, pattern: str) -> str | None:
        match = re.search(pattern, content, re.MULTILINE | re.IGNORECASE)
        if match:
            return match.group(1).strip().strip("*_")
        return None

    # ------------------------------------------------------------------
    # PDF export (raw PDF stream — for full branding use DOCX + LibreOffice)
    # ------------------------------------------------------------------

    def export_pdf(self, title: str, content: str, stem: str) -> str:
        """
        Basic text-only PDF export.
        For a fully branded PDF, open the exported DOCX in LibreOffice and
        use 'Export to PDF' for correct Candara fonts and styling.
        """
        target = Path(settings.export_dir) / f"{stem}.pdf"
        lines = [f"Course: {title}", ""] + content.splitlines()
        text = "\n".join(self._escape_pdf_text(line[:110]) for line in lines[:120])
        content_stream = f"BT /F1 11 Tf 50 790 Td 14 TL ({text}) Tj T* ET"
        content_bytes = content_stream.encode("latin-1", errors="replace")

        objects = [
            b"1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj\n",
            b"2 0 obj << /Type /Pages /Count 1 /Kids [3 0 R] >> endobj\n",
            b"3 0 obj << /Type /Page /Parent 2 0 R /MediaBox [0 0 595 842] /Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >> endobj\n",
            b"4 0 obj << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> endobj\n",
            f"5 0 obj << /Length {len(content_bytes)} >> stream\n".encode("latin-1")
            + content_bytes
            + b"\nendstream endobj\n",
        ]
        pdf = b"%PDF-1.4\n"
        offsets: list[int] = [0]
        for obj in objects:
            offsets.append(len(pdf))
            pdf += obj
        xref_offset = len(pdf)
        pdf += f"xref\n0 {len(offsets)}\n".encode("latin-1")
        pdf += b"0000000000 65535 f \n"
        for offset in offsets[1:]:
            pdf += f"{offset:010d} 00000 n \n".encode("latin-1")
        pdf += f"trailer << /Size {len(offsets)} /Root 1 0 R >>\nstartxref\n{xref_offset}\n%%EOF".encode("latin-1")
        target.write_bytes(pdf)
        return str(target)

    def _escape_pdf_text(self, value: str) -> str:
        return value.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)").replace("\n", "\\n")
